import requests
import json
from docx import Document
from docx.shared import Pt
import docx
from bs4 import BeautifulSoup
from pprint import pprint
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re



def save_and_download_image(picture_url,photo_path):
        response = requests.get(picture_url)
        with open(photo_path, "wb") as file:
            file.write(response.content)
        return response


user_choise = input("Write link to your wiki page: ")

if 'ru.wikipedia.org' not in user_choise:
    print("Try to write another link")
    exit()



response = requests.get(user_choise)
response.raise_for_status() 
soup = BeautifulSoup(response.text, features="html.parser")
all_tags = soup.find_all(['h1', 'h2','p','img'])
#for paragraph in all_tags:
    #print (paragraph.text)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
#head1 = doc.add_heading('Добавление заголовка документа', level=1)
#head2 = doc.add_heading('Основы работы с файлами Microsoft Word на Python.', level=2)
#doc.add_paragraph('майнкрафт')
#head1.alignment = WD_ALIGN_PARAGRAPH.CENTER
#head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
all_tags.reverse()
for tag in all_tags:
    if "<p" in str(tag) :
        #print(tag)
        #print(all_tags.index(tag))
        number = all_tags.index(tag)
        break
all_tags = all_tags[number:-2]
all_tags.reverse()
print(all_tags)


for tag in all_tags:
    sigmaetoybica = (re.sub(r'\[.*?\]',"",tag.text))
    if "<p>" in str(tag):
        paragraph = doc.add_paragraph(sigmaetoybica)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if "<img" in str(tag):
        url = f'https:{tag["src"]}'
        save_and_download_image(picture_url=url,photo_path="./image_papka.jpg")
        image = doc.add_picture('./image_papka.jpg')
        picture_paragraph = doc.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h2" in str(tag):
        head2 = doc.add_heading(sigmaetoybica, level=2)
        head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h1" in str(tag):
        head1 = doc.add_heading(sigmaetoybica, level=1)
        head1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h3" in str(tag):
        head3 = doc.add_heading(sigmaetoybica, level=3)
        head3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h4" in str(tag):
        head4 = doc.add_heading(sigmaetoybica, level=4)
        head4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h5" in str(tag):
        head5 = doc.add_heading(sigmaetoybica, level=5)
        head5.alignment = WD_ALIGN_PARAGRAPH.CENTER




doc.save('./test.docx')









